"""
Application Kit: a per-job, ready-to-use application package.

One click on a job page generates — grounded in the user's real resume —
  * a TAILORED RESUME (structured; downloadable as PDF and Word),
  * a complete COVER LETTER (PDF / copy),
  * drafted answers to the five standard portal questions,
all stored in `application_kits` so viewing and downloading never re-spends
AI calls. Regenerating overwrites in place. Generation costs 2 AI calls and
counts against the monthly AI budget (JOBBOT_LLM_MONTHLY_CAP).

Nothing is ever submitted anywhere: the kit produces files and text the user
pastes into the real application themselves.
"""

from __future__ import annotations

import io
import logging
import re
from datetime import date
from pathlib import Path

from fastapi import APIRouter, Depends, Request
from fastapi.responses import RedirectResponse, Response

from app.assist import LLMError, build_application_kit, tailor_resume_structured
from app.models import ApplicationKit, Resume, User
from app.web.deps import add_flash, get_db, render, require_user
from sqlalchemy.orm import Session as SessionType

logger = logging.getLogger("jobbot.web.kit")

router = APIRouter()

_DEJAVU = Path("/usr/share/fonts/truetype/dejavu")


# --------------------------------------------------------------- generation
@router.post("/jobs/{job_id}/kit")
def generate_kit(
    request: Request,
    job_id: int,
    user: User = Depends(require_user),
    db: SessionType = Depends(get_db),
):
    """HTMX: build (or rebuild) the kit and return the kit fragment."""
    from app.llm_budget import try_spend
    from app.web.assist import _job_for_user, _resume_text

    job = _job_for_user(db, user, job_id)
    if job is None:
        return render(request, "_kit.html", user=user,
                      kit_error="That job isn't in your matches.")
    resume_text = _resume_text(db, user)
    if not resume_text:
        return render(request, "_kit.html", user=user,
                      kit_error="Upload a resume first on the Resumes page.")
    if not try_spend(db, 2):  # two model calls per kit
        return render(request, "_kit.html", user=user,
                      kit_error="This month's AI budget is used up — the kit "
                                "can't be generated until it resets on the 1st.")

    try:
        content = build_application_kit(resume_text, job)
        content["resume"] = tailor_resume_structured(resume_text, job)
    except LLMError as exc:
        return render(request, "_kit.html", user=user, kit_error=str(exc))

    row = (
        db.query(ApplicationKit)
        .filter_by(user_id=user.id, job_id=job.id)
        .first()
    )
    if row is None:
        row = ApplicationKit(user_id=user.id, job_id=job.id)
        db.add(row)
    row.content = content
    db.commit()
    logger.info("built application kit for user %s job %s", user.id, job.id)
    return render(request, "_kit.html", user=user, job=job, kit=content, kit_row=row)


# ---------------------------------------------------------------- downloads
_ARTIFACTS = {
    "resume.pdf": ("resume", "pdf"),
    "resume.docx": ("resume", "docx"),
    "cover-letter.pdf": ("letter", "pdf"),
    "cover-letter.docx": ("letter", "docx"),
}
_MEDIA = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


@router.get("/jobs/{job_id}/kit/{fname}")
def download_artifact(
    request: Request,
    job_id: int,
    fname: str,
    user: User = Depends(require_user),
    db: SessionType = Depends(get_db),
):
    art = _ARTIFACTS.get(fname)
    row = (
        db.query(ApplicationKit)
        .filter_by(user_id=user.id, job_id=job_id)
        .first()
    )
    if art is None or row is None or not row.content:
        add_flash(request, "Generate the application kit first.", "error")
        return RedirectResponse(f"/jobs/{job_id}", status_code=303)

    kind, ext = art
    content = row.content
    resume = content.get("resume") or {}
    name = resume.get("name") or "resume"
    # Kit downloads come out in the look the user saved in the resume builder
    # (their newest resume that has one), falling back to Modern.
    look = (
        db.query(Resume.look)
        .filter(Resume.user_id == user.id, Resume.look.isnot(None))
        .order_by(Resume.uploaded_at.desc())
        .limit(1)
        .scalar()
    ) or "modern"
    if kind == "resume":
        blob = (_resume_pdf(resume, look=look) if ext == "pdf"
                else _resume_docx(resume, look=look))
        base = f"{name} - tailored resume"
    else:
        letter = content.get("cover_letter") or ""
        head = {"name": resume.get("name") or "", "contact": resume.get("contact") or ""}
        blob = (_letter_pdf(letter, look=look, **head) if ext == "pdf"
                else _letter_docx(letter, look=look, **head))
        base = f"{name} - cover letter"

    safe = re.sub(r"[^A-Za-z0-9 ._-]+", "", base).strip() or "document"
    return Response(
        content=blob,
        media_type=_MEDIA[ext],
        headers={"Content-Disposition": f'attachment; filename="{safe}.{ext}"'},
    )


# ------------------------------------------------------------- file builders
# Word twin of _PDF_LOOKS (defined below with the PDF builder).
_DOCX_LOOKS = {
    None:      {"font": "Calibri", "accent": None,            "center": False, "bar": False, "twocol": False, "underline": False},
    "modern":  {"font": "Calibri", "accent": (79, 70, 229),   "center": False, "bar": True,  "twocol": False, "underline": True},
    "classic": {"font": "Georgia", "accent": (25, 30, 48),    "center": True,  "bar": False, "twocol": False, "underline": True},
    "twocol":  {"font": "Calibri", "accent": (13, 148, 136),  "center": False, "bar": False, "twocol": True,  "underline": True},
    "minimal": {"font": "Calibri", "accent": (120, 128, 145), "center": False, "bar": False, "twocol": False, "underline": False},
}


def _docx_bottom_border(p, color: str, sz: int):
    """A bottom border on a paragraph — python-docx has no API for this."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    edge = OxmlElement("w:bottom")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(sz))  # eighths of a point
    edge.set(qn("w:space"), "2")
    edge.set(qn("w:color"), color)
    pbdr.append(edge)
    pPr.append(pbdr)


def _resume_docx(data: dict, look: str | None = None) -> bytes:
    """The tailored structured resume as a clean, ATS-safe Word document, styled
    to match the builder look when one is given. Sized with the same one-page
    scale the PDF settles on (Word can't measure its own pages)."""
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    st = _DOCX_LOOKS.get(look) or _DOCX_LOOKS[None]
    accent = st["accent"]
    k = _resume_fit_scale(data, look)
    d = docx.Document()
    style = d.styles["Normal"]
    style.font.name = st["font"]
    style.font.size = Pt(round(10.5 * k, 1))
    for s in d.sections:
        s.left_margin = s.right_margin = Inches(0.7)
        s.top_margin = s.bottom_margin = Inches(0.7)

    def para(text, *, into=None, size=10.5, bold=False, color=None, space_after=4):
        p = (into if into is not None else d).add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(round(size * k, 1))
        if color:
            run.font.color.rgb = RGBColor(*color)
        p.paragraph_format.space_after = Pt(round(space_after * k, 1))
        return p

    def heading(text, *, into=None):
        p = para((text or "").upper(), into=into, size=10.5, bold=True,
                 color=accent, space_after=2)
        if st["underline"]:
            _docx_bottom_border(p, "D5DAE4", 4)
        return p

    def section_block(sec, *, into=None):
        heading(sec.get("heading") or "", into=into)
        for ent in sec.get("entries") or []:
            head = " — ".join(x for x in [ent.get("role"), ent.get("org")] if x)
            line = " · ".join(x for x in [head, ent.get("dates")] if x)
            if line:
                para(line, into=into, bold=True, space_after=1)
            for b in ent.get("bullets") or []:
                p = (into if into is not None else d).add_paragraph(b, style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
        p = (into if into is not None else d).add_paragraph()
        p.paragraph_format.space_after = Pt(2)

    def skills_table(sec):
        """Skills in 2-3 horizontal columns (a borderless table, column-major),
        matching the on-screen preview. Main flow only — the two-column look's
        sidebar keeps its vertical list."""
        heading(sec.get("heading") or "")
        items = [b for ent in sec.get("entries") or []
                 for b in ent.get("bullets") or []]
        if not items:
            return
        ncols = 3 if len(items) >= 6 else 2 if len(items) >= 2 else 1
        nrows = -(-len(items) // ncols)
        t = d.add_table(rows=nrows, cols=ncols)
        t.allow_autofit = False
        for c in range(ncols):
            for r in range(nrows):
                cell = t.cell(r, c)
                cell.width = Inches(7.1 / ncols)
                idx = c * nrows + r
                if idx < len(items):
                    p = cell.paragraphs[0]
                    run = p.add_run(f"•  {items[idx]}")
                    run.font.size = Pt(round(10 * k, 1))
                    p.paragraph_format.space_after = Pt(1)
        spacer = d.add_paragraph()
        spacer.paragraph_format.space_after = Pt(2)

    # ---- header --------------------------------------------------------
    if st["bar"]:  # Modern's accent bar, as a full-width rule above the name
        p = d.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _docx_bottom_border(p, "4F46E5", 20)
    header = []
    if data.get("name"):
        header.append(para(data["name"], size=19, bold=True, space_after=1))
    if data.get("headline"):
        header.append(para(data["headline"], size=11,
                           color=accent or (70, 80, 105), space_after=1))
    if data.get("contact") and not st["twocol"]:
        header.append(para(data["contact"], size=9.5, color=(110, 120, 140),
                           space_after=8))
    if st["center"]:
        for p in header:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sections = data.get("sections") or []

    # ---- body: sidebar table (contact + skills) or one column -----------
    if st["twocol"]:
        table = d.add_table(rows=1, cols=2)
        table.allow_autofit = False
        side_cell, main_cell = table.rows[0].cells
        side_cell.width = Inches(2.3)
        main_cell.width = Inches(4.8)
        if data.get("contact"):
            heading("Contact", into=side_cell)
            para(data["contact"], into=side_cell, size=9, space_after=8)
        for sec in sections:
            if _is_skills(sec):
                section_block(sec, into=side_cell)
        if data.get("summary"):
            heading("Summary", into=main_cell)
            para(data["summary"], into=main_cell, space_after=8)
        for sec in sections:
            if not _is_skills(sec):
                section_block(sec, into=main_cell)
        for cell in (side_cell, main_cell):  # drop each cell's default blank first line
            first = cell.paragraphs[0]
            if len(cell.paragraphs) > 1 and not first.text:
                first._p.getparent().remove(first._p)
    else:
        if data.get("summary"):
            heading("Summary")
            para(data["summary"], space_after=8)
        for sec in sections:
            if _is_skills(sec):
                skills_table(sec)
            else:
                section_block(sec)

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _letter_docx(text: str, look: str | None = None,
                 name: str = "", contact: str = "") -> bytes:
    """The cover letter as a Word document — letterhead styled to the look."""
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    st = _DOCX_LOOKS.get(look) or _DOCX_LOOKS[None]
    d = docx.Document()
    d.styles["Normal"].font.name = st["font"]
    d.styles["Normal"].font.size = Pt(11)
    for s in d.sections:
        s.left_margin = s.right_margin = Inches(0.8)
        s.top_margin = s.bottom_margin = Inches(0.8)

    def para(text, *, size=11, bold=False, color=None, space_after=6):
        p = d.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        if look and st["center"]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space_after)
        return p

    if look:
        if st["bar"]:  # Modern's accent bar, as a full-width rule up top
            bar = d.add_paragraph()
            bar.paragraph_format.space_after = Pt(8)
            _docx_bottom_border(bar, "4F46E5", 20)
        if name:
            para(name, size=15, bold=True, space_after=1)
        if contact:
            p = para(contact, size=9, color=(110, 120, 140), space_after=10)
            accent_hex = "%02X%02X%02X" % (st["accent"] or (190, 196, 210))
            _docx_bottom_border(p, accent_hex, 8)
        today = date.today()
        p = para(f"{today.strftime('%B')} {today.day}, {today.year}",
                 size=10, color=(80, 90, 115), space_after=14)
        p.alignment = None  # the date stays left even in the centered classic look

    for part in (text or "").split("\n"):
        p = d.add_paragraph(part)
        p.paragraph_format.space_after = Pt(6 if part.strip() else 0)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _pdf_doc():
    """A5-margin A4 PDF with full-unicode DejaVu fonts registered."""
    from fpdf import FPDF

    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=16)
    pdf.set_margins(17, 16, 17)
    pdf.add_font("dj", "", str(_DEJAVU / "DejaVuSans.ttf"))
    pdf.add_font("dj", "B", str(_DEJAVU / "DejaVuSans-Bold.ttf"))
    pdf.add_font("djs", "", str(_DEJAVU / "DejaVuSerif.ttf"))
    pdf.add_font("djs", "B", str(_DEJAVU / "DejaVuSerif-Bold.ttf"))
    pdf.add_page()
    return pdf


# Print styling per builder look — the download twin of _resume_paper.html:
# font family, accent colour, header alignment, and the structural extras
# (Modern's accent bar, the two-column sidebar). look=None is the original
# neutral style, still used by the application-kit downloads.
_PDF_LOOKS = {
    None:      {"font": "dj",  "accent": (60, 66, 96),    "align": "L", "bar": False, "twocol": False, "rule": True,  "underline": False},
    "modern":  {"font": "dj",  "accent": (79, 70, 229),   "align": "L", "bar": True,  "twocol": False, "rule": False, "underline": True},
    "classic": {"font": "djs", "accent": (25, 30, 48),    "align": "C", "bar": False, "twocol": False, "rule": True,  "underline": True},
    "twocol":  {"font": "dj",  "accent": (13, 148, 136),  "align": "L", "bar": False, "twocol": True,  "rule": False, "underline": True},
    "minimal": {"font": "dj",  "accent": (120, 128, 145), "align": "L", "bar": False, "twocol": False, "rule": False, "underline": False},
}


def _is_skills(sec: dict) -> bool:
    return "skill" in (sec.get("heading") or "").lower()


# Shrink steps tried, in order, until the resume fits ONE page. 0.75 is the
# readability floor — a resume that still overflows there keeps its 2nd page.
_FIT_STEPS = (1.0, 0.95, 0.9, 0.85, 0.8, 0.75)


def _render_resume_pdf(data: dict, look: str | None, k: float):
    """Render the resume once at scale k (fonts, line heights and gaps all
    multiply by k) and return the FPDF object — the fit loop checks .page."""
    st = _PDF_LOOKS.get(look) or _PDF_LOOKS[None]
    fam, accent, align = st["font"], st["accent"], st["align"]
    pdf = _pdf_doc()

    def width():
        return pdf.w - pdf.l_margin - pdf.r_margin

    def mc(txt, *, size, h, style="", color, align="L", indent=0.0):
        """One text block, always starting at the left margin — fpdf leaves the
        cursor at the cell's right edge, which used to clip the next block."""
        pdf.set_font(fam, style, size * k)
        pdf.set_text_color(*color)
        pdf.set_x(pdf.l_margin + indent)
        pdf.multi_cell(width() - indent, h * k, txt, align=align)

    def heading(txt):
        mc((txt or "").upper(), size=10.2, h=5.6, style="B", color=accent)
        if st["underline"]:
            pdf.set_draw_color(205, 210, 222)
            pdf.line(pdf.l_margin, pdf.get_y() + 0.4,
                     pdf.l_margin + width(), pdf.get_y() + 0.4)
            pdf.ln(2.2 * k)
        else:
            pdf.ln(0.5 * k)

    def summary_block():
        heading("Summary")
        mc(data["summary"], size=9.6, h=5.0, color=(45, 52, 70))
        pdf.ln(2.5 * k)

    def section_block(sec, *, compact=False):
        heading(sec.get("heading") or "")
        for ent in sec.get("entries") or []:
            head = " — ".join(x for x in [ent.get("role"), ent.get("org")] if x)
            if head:
                mc(head, size=9.8, h=5.2, style="B", color=(30, 36, 58))
            if ent.get("dates"):
                mc(ent["dates"], size=8.6, h=4.6, color=(125, 132, 150))
            for b in ent.get("bullets") or []:
                mc(f"•  {b}", size=8.8 if compact else 9.4, h=4.9,
                   color=(45, 52, 70), indent=3)
            pdf.ln(1.2 * k)
        pdf.ln(1.4 * k)

    def skills_block(sec):
        """Skills flow into 2-3 horizontal columns (column-major), matching the
        on-screen preview — not one long vertical list."""
        heading(sec.get("heading") or "")
        items = [b for ent in sec.get("entries") or []
                 for b in ent.get("bullets") or []]
        if not items:
            pdf.ln(1.4 * k)
            return
        ncols = 3 if len(items) >= 6 else 2 if len(items) >= 2 else 1
        rows = -(-len(items) // ncols)
        col_w = width() / ncols
        y0 = pdf.get_y()
        y_end = y0
        pdf.set_font(fam, "", 9.4 * k)
        pdf.set_text_color(45, 52, 70)
        for c in range(ncols):
            pdf.set_y(y0)
            for item in items[c * rows:(c + 1) * rows]:
                pdf.set_x(pdf.l_margin + c * col_w)
                pdf.multi_cell(col_w - 4, 4.9 * k, f"•  {item}")
            y_end = max(y_end, pdf.get_y())
        pdf.set_y(y_end)
        pdf.ln(1.4 * k)

    # ---- header: bar (modern) / name / headline / contact -------------------
    if st["bar"]:
        pdf.set_fill_color(*accent)
        pdf.rect(pdf.l_margin, pdf.get_y(), 28, 2.4 * k, style="F")
        pdf.ln(6.5 * k)
    if data.get("name"):
        mc(data["name"], size=17, h=7.5, style="B", color=(20, 24, 40), align=align)
    if data.get("headline"):
        mc(data["headline"], size=10.5, h=5.4, align=align,
           color=(accent if look else (80, 90, 115)))
    if data.get("contact") and not st["twocol"]:
        mc(data["contact"], size=8.8, h=4.8, color=(120, 128, 145), align=align)
    if st["rule"]:
        pdf.ln(1.5 * k)
        pdf.set_draw_color(190, 196, 210)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln((3 if look != "minimal" else 5) * k)

    sections = data.get("sections") or []

    # ---- body: sidebar (contact + skills) beside the rest, or one column ----
    if st["twocol"]:
        l0, r0 = pdf.l_margin, pdf.r_margin
        side = (pdf.w - l0 - r0) * 0.32
        gap = 7
        top, page0 = pdf.get_y(), pdf.page

        pdf.set_right_margin(pdf.w - l0 - side)  # narrow to the sidebar
        if data.get("contact"):
            heading("Contact")
            mc(data["contact"], size=8.8, h=4.8, color=(45, 52, 70))
            pdf.ln(2.5 * k)
        for sec in sections:
            if _is_skills(sec):
                section_block(sec, compact=True)
        y_side = pdf.get_y()

        pdf.page = page0                          # main column starts back at the top
        pdf.set_left_margin(l0 + side + gap)
        pdf.set_right_margin(r0)
        pdf.set_y(top)
        if data.get("summary"):
            summary_block()
        for sec in sections:
            if not _is_skills(sec):
                section_block(sec)

        y_main = pdf.get_y() if pdf.page == page0 else pdf.page_break_trigger
        cur = pdf.page
        pdf.page = page0                          # divider between the columns
        pdf.set_draw_color(210, 215, 226)
        pdf.line(l0 + side + gap / 2, top, l0 + side + gap / 2, max(y_side, y_main))
        pdf.page = cur
        pdf.set_left_margin(l0)
    else:
        if data.get("summary"):
            summary_block()
        for sec in sections:
            if _is_skills(sec):
                skills_block(sec)
            else:
                section_block(sec)

    return pdf


def _resume_pdf(data: dict, look: str | None = None) -> bytes:
    for k in _FIT_STEPS:
        pdf = _render_resume_pdf(data, look, k)
        if pdf.page == 1:
            break
    return bytes(pdf.output())


def _resume_fit_scale(data: dict, look: str | None = None) -> float:
    """The scale at which this resume fits one PDF page — the Word builder uses
    it too (python-docx has no layout engine to measure pages with)."""
    for k in _FIT_STEPS:
        if _render_resume_pdf(data, look, k).page == 1:
            return k
    return _FIT_STEPS[-1]


def _letter_pdf(text: str, look: str | None = None,
                name: str = "", contact: str = "") -> bytes:
    """The cover letter as a PDF — with a letterhead (name, contact, date)
    styled to the given look when one is passed."""
    st = _PDF_LOOKS.get(look) or _PDF_LOOKS[None]
    fam, accent, align = st["font"], st["accent"], st["align"]
    pdf = _pdf_doc()
    w = pdf.w - pdf.l_margin - pdf.r_margin

    def mc(txt, *, size, h, style="", color, align="L"):
        # always restart at the left margin — fpdf leaves the cursor at the
        # right edge, which used to clip consecutive paragraphs off-page
        pdf.set_font(fam, style, size)
        pdf.set_text_color(*color)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(w, h, txt, align=align)

    if look:
        if st["bar"]:
            pdf.set_fill_color(*accent)
            pdf.rect(pdf.l_margin, pdf.get_y(), 28, 2.4, style="F")
            pdf.ln(6.5)
        if name:
            mc(name, size=15, h=6.8, style="B", color=(20, 24, 40), align=align)
        if contact:
            mc(contact, size=8.8, h=4.8, color=(120, 128, 145), align=align)
        pdf.ln(2)
        pdf.set_draw_color(*accent)  # the look's signature color as the divider
        pdf.set_line_width(0.5)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.set_line_width(0.2)
        pdf.ln(6)
        today = date.today()
        mc(f"{today.strftime('%B')} {today.day}, {today.year}",
           size=9.6, h=5.0, color=(80, 90, 115))
        pdf.ln(4)

    for part in (text or "").split("\n"):
        if part.strip():
            mc(part, size=10.6, h=5.6, color=(35, 42, 62))
        else:
            pdf.ln(3.4)
    return bytes(pdf.output())
